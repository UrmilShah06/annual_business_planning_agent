# planning_agent.py
# LangGraph 4-node workflow for annual business planning
# baseline -> target cascade -> loyalty -> word doc
# node 2 uses a proposer/challenger pattern with two GPT-4o calls

# Few Thoughts for next versions 
# The following will make planning agent smarter
# It will make the agent flexible for CEO, CFO, Planning Head, Business Heads to adopt, close the planning by staying in the agent itself
# - Go back 3 years of actuals instead of just 1 year — store and category trends
# - Output three plan versions — Conservative / Actual / Aggressive
#
# - Version the plan alongside the assumption file — if assumptions change so does the plan along with it
#   planning assumptions mid-cycle the old plan should still be retrievable,
#
# - Extend the human-in-the-loop review beyond store level
#   Currently only store growth override is possible, but in practice category managers also
#   push back on category splits and brand buyers negotiate OTB — need to
#   add category and brand level override capability
#
# - Add a store selector so the planning head can isolate one store and adjust
#   its category or brand mix without touching the rest of the chain plan
#
# - Pull in macro signals — consumption indices, Research reports, Government consumer indicator data —
#   use an LLM to extract inflation expectations, category trend signals
#   (ladies western growing faster in tier 2, menswear recovery post-pandemic),
#   HNI growth by city — feed this into the baseline assumptions so the plan
#   has an outside-in view, not just last year actuals


import os
import json
import re
from pathlib import Path
from typing import TypedDict, List, Optional, Dict
import pandas as pd
import numpy as np
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from langgraph.graph import StateGraph, END
from langgraph.checkpoint.memory import MemorySaver

DATA_DIR = "./data"
_cache = {}  # simple session cache - fine for single user, Redis would be better at scale


class PlanningState(TypedDict):
    company_growth_pct: float
    regional_growth: Dict[str, float]
    loyalty_target_pct: float
    loyalty_revenue_contrib_pct: float
    api_key: str
    baseline_analysis: str
    baseline_data: dict
    target_plan: str
    target_data: dict
    challenger_feedback: str
    loyalty_plan: str
    loyalty_data: dict
    human_adjustments: dict
    plan_approved: bool
    final_plan_text: str
    docx_path: str
    session_log: List[str]
    current_node: str
    error: Optional[str]


def load_df(source):
    # caching here because reloading from excel on every call was slow
    if source not in _cache:
        files = {
            "sales":    "01_FY25_Weekly_Sales_Actuals.xlsx",
            "brands":   "02_FY25_Brand_Performance.xlsx",
            "loyalty":  "03_FY25_Loyalty_Actuals.xlsx",
            "assumptions": "04_Planning_Assumptions.xlsx",
        }
        path = Path(DATA_DIR) / files[source]
        _cache[source] = pd.read_excel(path)
        _cache[source].columns = _cache[source].columns.str.strip()
    return _cache[source].copy()


def load_assumptions():
    path = Path(DATA_DIR) / "04_Planning_Assumptions.xlsx"
    store_df = pd.read_excel(path, sheet_name="Store Assumptions")
    brand_df = pd.read_excel(path, sheet_name="Brand Assumptions")
    params_df = pd.read_excel(path, sheet_name="Planning Parameters")
    # strip whitespace from column names - was causing key errors earlier
    store_df.columns = store_df.columns.str.strip()
    brand_df.columns = brand_df.columns.str.strip()
    params_df.columns = params_df.columns.str.strip()
    return {
        "stores": store_df,
        "brands": brand_df,
        "params": params_df,
    }


# langgraph was throwing serialization errors with numpy types
# had to write this conversion - not ideal but works
def _fix_types(obj):
    if isinstance(obj, dict):
        return {k: _fix_types(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [_fix_types(v) for v in obj]
    elif isinstance(obj, np.integer):
        return int(obj)
    elif isinstance(obj, np.floating):
        return float(obj)
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    return obj

# keeping old name as alias so I don't have to rename everywhere
to_python_types = _fix_types


# over-engineered on purpose — wraps round(v*100, 1) but felt safer having a named function
# after the LLM kept producing 0.4523 instead of 45.2 in early versions, wanted something explicit
def _safe_pct(v):
    try:
        return round(float(v) * 100, 1)
    except (TypeError, ValueError):
        return 0.0


# tried a version that also handled None and NaN here
# _safe_pct_v2 = lambda v: round(float(v)*100,1) if v is not None and not pd.isna(v) else 0.0
# scrapped it — lambda made the stack traces unreadable when things broke at 1am


def baseline_analysis_node(state):
    llm = ChatOpenAI(model="gpt-4o", temperature=0.1, openai_api_key=state["api_key"])

    sales_df = load_df("sales")
    brand_df = load_df("brands")

    store_totals = sales_df.groupby(["Store", "Region", "Store_Size", "Store_Maturity"]).agg(
        FY25_Revenue=("Actual_INR_L", "sum"),
        FY25_Target=("Target_INR_L", "sum"),
        Total_Bills=("Bills_Count", "sum"),
        Avg_Conversion=("Conversion_%", "mean"),
        Avg_UPT=("UPT", "mean"),
    ).reset_index()
    store_totals["Achievement_%"] = (store_totals["FY25_Revenue"] / store_totals["FY25_Target"] * 100).round(1)
    store_totals["FY25_Revenue"] = store_totals["FY25_Revenue"].round(2)

    # category mix - need this for the cascade later
    cat_mix = sales_df.groupby(["Store", "Category"]).agg(
        Revenue=("Actual_INR_L", "sum")
    ).reset_index()

    weekly_chain = sales_df.groupby("Week_Number")["Actual_INR_L"].sum().reset_index()
    peak_weeks = weekly_chain.nlargest(5, "Actual_INR_L")["Week_Number"].tolist()

    # groupby is cleaner here
    brand_summary = brand_df.groupby(["Category", "Brand"]).agg(
        Avg_ST=("Sell_Through_%", "mean"),
        Total_Revenue=("Revenue_INR_L", "sum"),
    ).reset_index().sort_values("Avg_ST", ascending=False)

    baseline_data = {}
    for _, row in store_totals.iterrows():
        baseline_data[row["Store"]] = {
            "region": row["Region"],
            "size": row["Store_Size"],
            "maturity": row["Store_Maturity"],
            "fy25_revenue": float(row["FY25_Revenue"]),
            "achievement_pct": float(row["Achievement_%"]),
            "total_bills": int(row["Total_Bills"]),
            "avg_conversion": float(row["Avg_Conversion"]),
            "categories": {}
        }
        store_cats = cat_mix[cat_mix["Store"] == row["Store"]]
        total_rev = store_cats["Revenue"].sum()
        for _, crow in store_cats.iterrows():
            baseline_data[row["Store"]]["categories"][crow["Category"]] = {
                "revenue": float(round(crow["Revenue"], 2)),
                "share_pct": float(round(crow["Revenue"] / total_rev * 100, 1))
            }

    summary_lines = []
    for store, data in baseline_data.items():
        summary_lines.append(
            f"{store} ({data['region']}, {data['size']}, {data['maturity']}): "
            f"FY25 Revenue INR {data['fy25_revenue']}L, "
            f"Achievement {data['achievement_pct']}%, "
            f"Bills {data['total_bills']:,}"
        )

    chain_total = sum(d["fy25_revenue"] for d in baseline_data.values())

    # tried sorting brand_summary by revenue instead of ST here — the LLM kept flagging
    # high-revenue low-ST brands as "strategic priorities" which was wrong
    # ST is the right signal for planning — a brand doing 40% ST needs space reduction not more OTB
    # brand_summary_by_rev = brand_summary.sort_values("Total_Revenue", ascending=False)

    top_brands = brand_summary.head(10)
    brand_lines = [f"{r['Brand']} ({r['Category']}): ST {r['Avg_ST']:.1f}%, Rev {r['Total_Revenue']:.1f}L"
                   for _, r in top_brands.iterrows()]

    response = llm.invoke([
        SystemMessage(content="""You are a Senior Retail Business Planning Analyst.
You specialise in building annual plans for large format retail chains.
Analyse FY25 performance data and produce a clear baseline summary for planning."""),
        HumanMessage(content=f"""Analyse FY25 performance and summarise the planning baseline.

CHAIN SUMMARY:
Total FY25 Revenue: INR {chain_total:.1f}L across 10 stores
Peak selling weeks: {peak_weeks}

STORE PERFORMANCE:
{chr(10).join(summary_lines)}

TOP PERFORMING BRANDS (by sell-through):
{chr(10).join(brand_lines)}

Provide:
1. CHAIN PERFORMANCE HEADLINE: Key FY25 metrics and overall health
2. REGIONAL ANALYSIS: North/South/East/West revenue contribution and growth trends
3. CATEGORY ANALYSIS: Which categories are growing, which are declining
4. STORE TIER ANALYSIS: Flagship vs standard vs new store performance gaps
5. SEASONAL INSIGHTS: Which weeks drove disproportionate revenue (peak weeks)
6. PLANNING IMPLICATIONS: Key factors that must be reflected in FY26 targets

Be specific with store names, INR values, and percentages.""")
    ])

    state["baseline_analysis"] = response.content
    state["baseline_data"] = to_python_types(baseline_data)
    state["session_log"].append(f"Agent 1 (Baseline): FY25 analysed — Chain INR {chain_total:.1f}L")
    state["current_node"] = "baseline_complete"
    return state


def target_cascade_node(state):
    # two GPT-4o calls - proposer sets targets, challenger stress tests
    # more robust than asking one model to do both - it tends to be too lenient on its own outputs
    llm = ChatOpenAI(model="gpt-4o", temperature=0.1, openai_api_key=state["api_key"])

    baseline_data = state["baseline_data"]
    company_growth = state["company_growth_pct"] / 100
    regional_growth = {k: v/100 for k, v in state["regional_growth"].items()}
    assumptions = load_assumptions()
    store_assumptions = assumptions["stores"]
    brand_assumptions = assumptions["brands"]

    # maturity caps - went back and forth on these numbers
    # 13% for mature is conservative but these stores have limited upside
    growth_caps = {"Mature": 0.13, "Mid": 0.20, "New": 0.40}

    chain_fy25 = sum(d["fy25_revenue"] for d in baseline_data.values())
    chain_target = chain_fy25 * (1 + company_growth)

    # chain_target computed but not passed to proposer — LLM was hallucinating chain numbers
    # when given both store-level and chain-level targets simultaneously, summed wrong ~30% of the time
    # safer to let it derive chain from store sum itself

    target_data = {}
    proposal_lines = []

    for store, data in baseline_data.items():
        region = data["region"]
        maturity = data["maturity"]
        fy25_rev = data["fy25_revenue"]

        store_row = store_assumptions[store_assumptions["Store"] == store]
        override = None
        if not store_row.empty:
            ov = store_row["Store_Growth_Override_%"].values[0]
            if pd.notna(ov) and str(ov) != "":
                try:
                    override = float(ov) / 100
                except:
                    pass  # if it's not a number just skip it

        floor_change = 0
        if not store_row.empty:
            fc = store_row["FY26_Floor_Change_%"].values[0]
            if pd.notna(fc):
                floor_change = float(fc) / 100

        # food court and SIS bonuses
        # TODO: make these configurable from planning assumptions sheet instead of hardcoding
        # each store will have a different multiplier based on area of food court and brands coming in SIS
        fc_bonus = 0
        if not store_row.empty:
            if store_row["New_Food_Court"].values[0] == "Y":
                fc_bonus += 0.05
            if store_row["New_SIS"].values[0] == "Y":
                fc_bonus += 0.03

        if override:
            growth_rate = override
        else:
            reg_growth = regional_growth.get(region, company_growth)
            cap = growth_caps[maturity]
            # underperformers shouldn't get the same target as stores that hit plan
            # this adjustment moderates their targets slightly
            ach_adj = (data["achievement_pct"] - 100) / 100 * 0.3
            growth_rate = min(cap, reg_growth + ach_adj + floor_change + fc_bonus)
            growth_rate = max(0.05, growth_rate)

        fy26_target = round(fy25_rev * (1 + growth_rate), 2)

        category_targets = {}
        for cat, cdata in data["categories"].items():
            cat_share = cdata["share_pct"] / 100
            category_targets[cat] = round(fy26_target * cat_share, 2)

        brand_df = load_df("brands")
        brand_otb = {}

        for cat, cat_target in category_targets.items():
            cat_brands = brand_df[(brand_df["Category"] == cat)].groupby("Brand").agg(
                Revenue=("Revenue_INR_L", "sum"),
                Avg_ST=("Sell_Through_%", "mean"),
                Avg_MRP=("MRP_INR", "mean"),
                Avg_Margin=("Margin_%", "mean"),
            ).reset_index()

            if cat_brands.empty:
                continue

            total_brand_rev = cat_brands["Revenue"].sum()
            sell_through_target = 0.72  # 72% ST is our planning assumption

            for _, brand_row in cat_brands.iterrows():
                brand = brand_row["Brand"]
                brand_share = brand_row["Revenue"] / total_brand_rev if total_brand_rev > 0 else 0

                b_row = brand_assumptions[
                    (brand_assumptions["Brand"] == brand) &
                    (brand_assumptions["Category"] == cat)
                ]
                space_mult = 1.0
                if not b_row.empty:
                    space_mult = float(b_row["OTB_Multiplier"].values[0])

                brand_revenue_target = cat_target * brand_share * space_mult
                mrp = brand_row["Avg_MRP"] if brand_row["Avg_MRP"] > 0 else 1299
                margin_pct = brand_row["Avg_Margin"] / 100 if brand_row["Avg_Margin"] > 0 else 0.52
                cost_pct = 1 - margin_pct

                # OTB formula: Revenue Target / (MRP * Sell Through) gives units
                # OTB value = otb units * mrp * cost price
                otb_units = int(brand_revenue_target * 100000 / (mrp * sell_through_target))
                
                otb_value = round(otb_units * mrp * cost_pct / 100000, 2)

                brand_otb[f"{cat}|{brand}"] = {
                    "category": cat,
                    "brand": brand,
                    "revenue_target_L": round(brand_revenue_target, 2),
                    "otb_units": otb_units,
                    "otb_value_L": otb_value,
                    "space_multiplier": space_mult,
                }

        target_data[store] = {
            "region": region,
            "maturity": maturity,
            "fy25_revenue": fy25_rev,
            "fy26_target": fy26_target,
            "growth_rate_pct": round(growth_rate * 100, 1),
            "category_targets": category_targets,
            "brand_otb": brand_otb,
        }

        proposal_lines.append(
            f"{store}: FY25 INR {fy25_rev}L → FY26 Target INR {fy26_target}L "
            f"(+{growth_rate*100:.1f}% growth, {maturity} store, {region})"
        )

    chain_fy26 = sum(d["fy26_target"] for d in target_data.values())

    proposer_response = llm.invoke([
        SystemMessage(content="""You are a Retail Planning Director proposing FY26 targets.
Present the target cascade clearly with business rationale for each store's growth rate."""),
        HumanMessage(content=f"""Present the FY26 target plan.

COMPANY TARGET: INR {chain_fy26:.1f}L (growth of {company_growth*100:.1f}% on FY25 INR {chain_fy25:.1f}L)

STORE TARGETS:
{chr(10).join(proposal_lines)}

Format as:
**COMPANY TARGET SUMMARY**
Chain revenue target and overall growth rationale.

**REGIONAL BREAKDOWN**
For each region use exactly this format:
Region: FY26 Target INR XXX.XL | Avg Growth XX.X% | Store count and mix (Mature/Mid/New)

**STORE-BY-STORE TARGETS**
For EACH store use exactly this format — do not skip the numbers:
Store_Name: FY26 Target INR XXX.XL (+XX.X% growth) | Maturity: XXX | Rationale: one sentence

**CATEGORY HIGHLIGHTS**
Top 3 categories by absolute target growth. Use ONLY these category names: Ladies Western, Ladies Ethnic, Men Formal, Men Casual, Kids, Sportswear, Footwear. Do NOT invent categories like Electronics or Home & Living.""")
 
    ])

    # challenger - CFO has no motivation to defend the plan
    # this is the check that catches targets that look good on paper but won't hold up in review
    challenger_response = llm.invoke([
        SystemMessage(content="""You are a Retail CFO acting as a planning challenger.
Your job is to stress-test proposed targets for realism and flag risks.
Be constructive but rigorous — identify targets that may be too aggressive or too conservative."""),
        HumanMessage(content=f"""Challenge these FY26 proposed targets.

PROPOSED TARGETS:
{chr(10).join(proposal_lines)}

CONTEXT:
- New stores (Pune_Kothrud, Hyderabad_Banjara) are expected to grow 35-38%
- Mature stores (Mumbai_Thane, Delhi_VasantKunj) capped at 12-13%
- Company total target: INR {chain_fy26:.1f}L

For each store, assess:
1. Is the growth rate realistic given maturity and market conditions?
2. Are there execution risks (floor plate changes, new categories) that could miss targets?
3. Are any targets too conservative (leaving growth on the table)?
4. What are the top 3 risks to the plan overall?

Format as:
**CHALLENGED TARGETS** (flag each as REALISTIC / AGGRESSIVE / CONSERVATIVE)
**TOP 3 PLAN RISKS**
**RECOMMENDED ADJUSTMENTS** (specific stores/rates to reconsider)""")
    ])

    state["target_plan"] = proposer_response.content
    state["target_data"] = to_python_types(target_data)
    state["challenger_feedback"] = challenger_response.content
    state["session_log"].append(
        f"Agent 2 (Targets): Chain FY26 target INR {chain_fy26:.1f}L "
        f"({company_growth*100:.1f}% growth). Challenger feedback generated."
    )
    state["current_node"] = "targets_complete"
    return state


def loyalty_planning_node(state):
    llm = ChatOpenAI(model="gpt-4o", temperature=0.1, openai_api_key=state["api_key"])

    loyalty_df = load_df("loyalty")
    target_data = state["target_data"]
    loyalty_target = state["loyalty_target_pct"] / 100
    loyalty_contrib = state["loyalty_revenue_contrib_pct"] / 100

    store_loyalty = loyalty_df.groupby("Store").agg(
        FY25_Total_Bills=("Total_Bills", "sum"),
        FY25_Loyalty_Bills=("Loyalty_Bills", "sum"),
        FY25_Total_Revenue=("Total_Revenue_INR_L", "sum"),
        FY25_Loyalty_Revenue=("Loyalty_Revenue_INR_L", "sum"),
        FY25_New_Enrollments=("New_Enrollments", "sum"),
        FY25_Active_Members=("Active_Members", "mean"),
        Avg_Basket_Loyalty=("Avg_Basket_Loyalty_INR", "mean"),
        Avg_Basket_NonLoyalty=("Avg_Basket_NonLoyalty_INR", "mean"),
    ).reset_index()

    loyalty_data = {}
    loyalty_lines = []

    for _, row in store_loyalty.iterrows():
        store = row["Store"]
        fy25_pen = row["FY25_Loyalty_Bills"] / row["FY25_Total_Bills"] if row["FY25_Total_Bills"] > 0 else 0

        # NOTE: Peak_Visits from IoT data would improve this — high footfall low penetration stores
        # need different activation vs low footfall high penetration stores
        # didn't pass it here because the agent started fixating on staffing recommendations instead of loyalty tactics
        # leaving it out for now, the bill-based penetration is sufficient for planning purposes

        # stretch factor - stores already above 50% penetration don't need as aggressive a push
        # learned this the hard way in planning cycles - setting same stretch for all stores
        # leads to sandbagging from high performers and unrealistic targets for laggards
        if fy25_pen > 0.50:
            stretch = 0.03
        elif fy25_pen > 0.35:
            stretch = 0.06
        else:
            stretch = 0.10

        # ops teams push back beyond 72% — called it out in the plan as a hard ceiling
        # tried 75% once, got told it was "aspirational not operational" in the review
        fy26_pen_target = min(0.72, fy25_pen + stretch)

        fy26_revenue = target_data.get(store, {}).get("fy26_target", row["FY25_Total_Revenue"] * 1.15)
        loyalty_rev_target = round(fy26_revenue * loyalty_contrib, 2)

        current_members = int(row["FY25_Active_Members"])
        target_members = int(current_members * (1 + stretch * 2))
        enrollment_target = int((target_members - current_members) * 1.25)  # 25% buffer for churn

        loyalty_data[store] = {
            "fy25_penetration_pct": round(fy25_pen * 100, 1),
            "fy26_penetration_target_pct": round(fy26_pen_target * 100, 1),
            "fy26_loyalty_revenue_target_L": loyalty_rev_target,
            "fy26_total_revenue_target_L": fy26_revenue,
            "loyalty_contribution_pct": round(loyalty_contrib * 100, 1),
            "new_enrollment_target": enrollment_target,
            "current_active_members": current_members,
            "avg_basket_premium_INR": int(row["Avg_Basket_Loyalty"]) - int(row["Avg_Basket_NonLoyalty"]),
        }

        loyalty_lines.append(
            f"{store}: FY25 penetration {fy25_pen*100:.1f}% → "
            f"FY26 target {fy26_pen_target*100:.1f}% | "
            f"Loyalty Revenue Target INR {loyalty_rev_target}L | "
            f"New Enrollments: {enrollment_target:,}"
        )

    chain_loyalty_rev = sum(d["fy26_loyalty_revenue_target_L"] for d in loyalty_data.values())
    chain_enrollments = sum(d["new_enrollment_target"] for d in loyalty_data.values())

    response = llm.invoke([
        SystemMessage(content="""You are a Retail Loyalty Programme Planning Specialist.
You build loyalty targets that are stretching but achievable, with clear
segment-level strategies for acquisition, retention, and reactivation."""),
        HumanMessage(content=f"""Build the FY26 loyalty plan.

CHAIN LOYALTY TARGETS:
Chain penetration target: {loyalty_target*100:.1f}%
Loyalty revenue contribution: {loyalty_contrib*100:.1f}% of total revenue
Chain loyalty revenue target: INR {chain_loyalty_rev:.1f}L
Total new enrollment target: {chain_enrollments:,}

STORE-LEVEL TARGETS:
{chr(10).join(loyalty_lines)}

Produce:
**CHAIN LOYALTY HEADLINE**
Penetration target, revenue target, enrollment target with growth %.

**STORE-LEVEL LOYALTY TARGETS**
For each store: FY25 penetration, FY26 target, stretch, rationale.

**SEGMENT STRATEGY**
- Premium segment (top 20% spenders): retention focus, exclusive benefits
- Regular segment (middle 50%): upgrade to premium path
- Occasional segment (bottom 30%): activation campaigns, milestone rewards

**KEY LOYALTY INITIATIVES FOR FY26**
3-4 specific programmes to drive enrollment and penetration.""")
    ])

    state["loyalty_plan"] = response.content
    state["loyalty_data"] = to_python_types(loyalty_data)
    state["session_log"].append(
        f"Agent 3 (Loyalty): Chain loyalty revenue target INR {chain_loyalty_rev:.1f}L, "
        f"Enrollments {chain_enrollments:,}"
    )
    state["current_node"] = "loyalty_complete"
    return state


def business_plan_writer_node(state):
    # final node - pulls everything together into a word document
    # this runs after human approval so we have the adjusted targets here
    # kept this node simple on purpose - extra LLM calls to format sections weren't worth the latency
    llm = ChatOpenAI(model="gpt-4o", temperature=0.1, openai_api_key=state["api_key"])

    target_data = state["target_data"]
    loyalty_data = state["loyalty_data"]
    company_growth = state["company_growth_pct"]

    chain_fy25 = sum(d["fy25_revenue"] for d in target_data.values())
    chain_fy26 = sum(d["fy26_target"] for d in target_data.values())
    chain_loyalty = sum(d["fy26_loyalty_revenue_target_L"] for d in loyalty_data.values())

    # chain totals passed explicitly because LLM summed wrong about 30% of the time
    # when asked to derive from store-level data in the same prompt
    response = llm.invoke([
        SystemMessage(content="""You are a Chief Strategy Officer writing the annual business plan.
This document will be presented to the Board. Language should be clear, strategic, and data-backed.
Use retail terminology correctly: OTB, sell-through, penetration, LFL (like-for-like), ATV."""),
        HumanMessage(content=f"""Write the FY26 Annual Business Plan narrative.

PLAN NUMBERS:
FY25 Actual Chain Revenue: INR {chain_fy25:.1f}L
FY26 Target Chain Revenue: INR {chain_fy26:.1f}L
Growth: {company_growth:.1f}%
FY26 Loyalty Revenue Target: INR {chain_loyalty:.1f}L
Chain Loyalty Penetration Target: {state.get('loyalty_target_pct', 45) if isinstance(state, dict) else getattr(state, 'loyalty_target_pct', 45)}%

PERFORMANCE BASELINE (from Agent 1):
{state.get('baseline_analysis', '') if isinstance(state, dict) else getattr(state, 'baseline_analysis', '')[:600]}

TARGET PLAN (from Agent 2):
{state.get('target_plan', '') if isinstance(state, dict) else getattr(state, 'target_plan', '')[:600]}

CHALLENGER FEEDBACK (from AutoGen challenger):
{state.get('challenger_feedback', '') if isinstance(state, dict) else getattr(state, 'challenger_feedback', '')[:400]}

LOYALTY PLAN (from Agent 3):
{state.get('loyalty_plan', '') if isinstance(state, dict) else getattr(state, 'loyalty_plan', '')[:600]}

Write:
**EXECUTIVE SUMMARY** (200 words)
FY26 ambition, key growth drivers, strategic priorities.

**FY26 TARGETS AT A GLANCE**
Chain revenue, growth %, loyalty revenue, penetration target.

**REGIONAL STRATEGY**
North / South / East / West targets and strategic focus per region.

**STORE PERFORMANCE TARGETS**
Table format: Store | FY25 Revenue | FY26 Target | Growth% | Key Initiative

**CATEGORY STRATEGY** — write one paragraph each for ONLY these 7 categories: Ladies Western, Ladies Ethnic, Men Formal, Men Casual, Kids, Sportswear, Footwear. Do NOT use any other category names.

**LOYALTY PROGRAMME PLAN**
Penetration targets, enrollment targets, 3 key initiatives.

**KEY RISKS AND MITIGATIONS**
Top 3 risks (from challenger) with mitigation strategies.

**PLANNING ASSUMPTIONS**
Growth rate, regional indices, loyalty contribution, OTB sell-through target.""")
    ])

    plan_text = response.content
    state["final_plan_text"] = plan_text

    docx_path = generate_word_document(state, plan_text)
    state["docx_path"] = docx_path

    state["session_log"].append(
        f"Agent 4 (Plan Writer): Business plan generated. "
        f"FY26 Chain Target INR {chain_fy26:.1f}L. Document saved."
    )
    state["current_node"] = "plan_complete"
    return state


def generate_word_document(state, plan_text):
    # generates the Word document from plan text + structured data tables
    # had to handle ImportError because python-docx isn't always installed
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title = doc.add_heading("FY26 ANNUAL BUSINESS PLAN", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.runs[0]
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        subtitle = doc.add_paragraph("Retail Chain — Strategic Growth Plan")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        subtitle.runs[0].font.size = Pt(11)

        growth_pct = state.get('company_growth_pct', 15) if isinstance(state, dict) else getattr(state, 'company_growth_pct', 15)
        doc.add_paragraph(f"Prepared by: Annual Business Planning Agent | "
                          f"Company Growth Target: {growth_pct:.1f}%")

        doc.add_paragraph("─" * 80)

        # tried using python-docx styles more aggressively here — custom styles, table of contents
        # abandoned it, the default styles vary too much across Word versions
        # board recipients were seeing different formatting on Mac vs Windows, not worth it
        for line in plan_text.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue
            if line.startswith('**') and line.endswith('**'):
                h = doc.add_heading(line.strip('*'), level=2)
                h.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            elif line.startswith('- ') or line.startswith('• '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.'):
                doc.add_paragraph(line, style='List Number')
            else:
                doc.add_paragraph(line)

        doc.add_heading("STORE TARGETS SUMMARY TABLE", level=2)
        target_data = state["target_data"]
        loyalty_data = state["loyalty_data"]

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        headers = ["Store", "Region", "FY25 Revenue (L)", "FY26 Target (L)", "Growth %", "Loyalty Target %"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].font.bold = True

        for store, data in target_data.items():
            row = table.add_row().cells
            row[0].text = store
            row[1].text = data["region"]
            row[2].text = f"INR {data['fy25_revenue']:.1f}L"
            row[3].text = f"INR {data['fy26_target']:.1f}L"
            row[4].text = f"+{data['growth_rate_pct']:.1f}%"
            loy = loyalty_data.get(store, {})
            row[5].text = f"{loy.get('fy26_penetration_target_pct', 45):.1f}%"

        doc.add_heading("BRAND OTB SUMMARY (ALL BRANDS)", level=2)

        all_otb = []
        for store, data in target_data.items():
            for key, otb in data.get("brand_otb", {}).items():
                all_otb.append({
                    "Store": store,
                    "Category": otb["category"],
                    "Brand": otb["brand"],
                    "Revenue Target (L)": otb["revenue_target_L"],
                    "OTB Units": otb["otb_units"],
                    "OTB Value (L)": otb["otb_value_L"],
                })

        otb_df = pd.DataFrame(all_otb)
        if not otb_df.empty:
            # aggregating at chain level for the Word doc - store level was too granular for the Board
            chain_otb = otb_df.groupby(["Category", "Brand"]).agg(
                Rev_Target=("Revenue Target (L)", "sum"),
                OTB_Units=("OTB Units", "sum"),
                OTB_Value=("OTB Value (L)", "sum"),
            ).reset_index().sort_values("OTB_Value", ascending=False)

            otb_table = doc.add_table(rows=1, cols=5)
            otb_table.style = 'Table Grid'
            ohdr = otb_table.rows[0].cells
            for i, h in enumerate(["Category", "Brand", "Rev Target (L)", "OTB Units", "OTB Value (L)"]):
                ohdr[i].text = h
                ohdr[i].paragraphs[0].runs[0].font.bold = True

            for _, row in chain_otb.iterrows():
                r = otb_table.add_row().cells
                r[0].text = row["Category"]
                r[1].text = row["Brand"]
                r[2].text = f"INR {row['Rev_Target']:.1f}L"
                r[3].text = f"{int(row['OTB_Units']):,}"
                r[4].text = f"INR {row['OTB_Value']:.1f}L"

        out_path = "./data/FY26_Annual_Business_Plan.docx"
        doc.save(out_path)
        return out_path

    except ImportError:
        # fallback if python-docx not installed
        out_path = "./data/FY26_Annual_Business_Plan.txt"
        with open(out_path, 'w') as f:
            f.write(plan_text)
        return out_path


def build_planning_workflow():
    workflow = StateGraph(PlanningState)
    workflow.add_node("baseline_analysis_node", baseline_analysis_node)
    workflow.add_node("target_cascade_node", target_cascade_node)
    workflow.add_node("loyalty_planning_node", loyalty_planning_node)
    workflow.add_node("business_plan_writer_node", business_plan_writer_node)
    workflow.set_entry_point("baseline_analysis_node")
    workflow.add_edge("baseline_analysis_node", "target_cascade_node")
    workflow.add_edge("target_cascade_node", "loyalty_planning_node")
    workflow.add_edge("loyalty_planning_node", "business_plan_writer_node")
    workflow.add_edge("business_plan_writer_node", END)
    # tried parallel execution for nodes 1 and 2 here — loyalty needs target data so couldn't parallelize fully
    # sequential is cleaner anyway, easier to debug when the LLM goes off script
    memory = MemorySaver()
    return workflow.compile(
        checkpointer=memory,
        interrupt_before=["business_plan_writer_node"]
    )


def get_initial_planning_state(
    company_growth_pct,
    regional_growth,
    loyalty_target_pct,
    loyalty_revenue_contrib_pct,
    api_key,
):
    return PlanningState(
        company_growth_pct=company_growth_pct,
        regional_growth=regional_growth,
        loyalty_target_pct=loyalty_target_pct,
        loyalty_revenue_contrib_pct=loyalty_revenue_contrib_pct,
        api_key=api_key,
        baseline_analysis="",
        baseline_data={},
        target_plan="",
        target_data={},
        challenger_feedback="",
        loyalty_plan="",
        loyalty_data={},
        human_adjustments={},
        plan_approved=False,
        final_plan_text="",
        docx_path="",
        session_log=[],
        current_node="start",
        error=None,
    )
